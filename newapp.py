from nicer import *
from db import *
from py_uis.selectReportView import Ui_selectReportView
from py_uis.curTableView import Ui_curTableView
from py_uis.tablesView import Ui_tableViewWindow
from py_uis.addRowView import Ui_addRowView
from PySide6.QtWidgets import *
from PySide6.QtCore import *
from PySide6 import QtCore
from PySide6.QtGui import *
import pandas as pd
from docx import Document

class TableViewWindow(QMainWindow):
    def __init__(self):
        super(TableViewWindow, self).__init__()
        self.ui = Ui_tableViewWindow()
        self.ui.setupUi(self)

        self.tablesList = self.__requestTables()
        for i in self.tablesList:
            listItem = QListWidgetItem(niceTableNames[i])
            listItem.setData(1,i)
            self.ui.listWidget.addItem(listItem)
        self.connect(self.ui.listWidget,
                    QtCore.SIGNAL("itemDoubleClicked(QListWidgetItem *)"),
                    self.goToTable)

        self.ui.reportBtn.triggered.connect(self.startReport)
        self.updateStatus("ОК")
    
    def startReport(self):
        print("Starting report")
        self.dialog = SelectReportDialog(self)
        self.dialog.show()
        self.dialog.rejected.connect(lambda: self.show())
        self.hide()

    def goToTable(self, listItem:QListWidgetItem):
        print(f"Table {listItem.data(1)} chosen")
        self.newWindow = CurTableViewWindow(listItem.data(1),self)
        self.hide()
        self.newWindow.show()

    def __requestTables(self) -> list:
        return list(i[0] for i in DBParser.execute("select table_name from information_schema.tables where table_schema = 'public' and table_type='BASE TABLE'"))
    
    def updateStatus(self, newStatus:str):
        self.ui.statusbar.showMessage(newStatus)

class AddRowDialog(QDialog):
    def __init__(self, colNames:list, tableName:str):
        super(AddRowDialog, self).__init__()
        self.ui = Ui_addRowView()
        self.ui.setupUi(self)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.setSizePolicy(QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Minimum)
        self.cols = {i:QLineEdit(self) for i in colNames}

        for colInd in range(len(self.cols.keys())):
            colName = list(self.cols.keys())[colInd]
            label = QLabel(niceTableHeaders[tableName][colName], self)
            self.ui.formLayout.setWidget(colInd, QFormLayout.ItemRole.LabelRole, label)
            self.ui.formLayout.setWidget(colInd, QFormLayout.ItemRole.FieldRole, self.cols[colName])

    def closeEvent(self, arg__1):
        self.setResult(0)
        return super().closeEvent(arg__1)

class EditRowDialog(QDialog):
    def __init__(self, colNames:dict, tableName:str):
        super(EditRowDialog, self).__init__()
        self.ui = Ui_addRowView()
        self.ui.setupUi(self)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.setSizePolicy(QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Minimum)
        self.cols = {k:QLineEdit(self,text=v) for k,v in colNames.items()}

        for colInd in range(len(self.cols.keys())):
            colName = list(self.cols.keys())[colInd]
            label = QLabel(niceTableHeaders[tableName][colName], self)
            self.ui.formLayout.setWidget(colInd, QFormLayout.ItemRole.LabelRole, label)
            self.ui.formLayout.setWidget(colInd, QFormLayout.ItemRole.FieldRole, self.cols[colName])

    def closeEvent(self, arg__1):
        self.setResult(0)
        return super().closeEvent(arg__1)

class CurTableViewWindow(QMainWindow):
    def __init__(self, tableName:str, parw:TableViewWindow):
        super(CurTableViewWindow, self).__init__()
        self.ui = Ui_curTableView()
        self.ui.setupUi(self)

        self.ACTIVE = True

        self.tableName = tableName
        self.ui.tableNameLabel.setText(niceTableNames[tableName])

        self.parw = parw

        self.curHeaders = list(map(lambda x: x[0],DBParser.execute(f"select column_name from information_schema.columns where table_schema = 'public' and table_name = '{self.tableName}';")))
        print(self.curHeaders)
        self.ui.tableWidget.verticalHeader().setVisible(False)

        class LineEdit(QLineEdit):
            def init(self,arg_1,parw):
                self.setPlaceholderText(arg_1)
                self.parw = parw
                self.eventHandled = False
                return self
            def focusOutEvent(self, arg__1):
                print("FILTER CELL LOST FOCUS")
                if not self.eventHandled:
                    self.eventHandled = True
                    print("FILTER CELL DONT HAVE EVENT_HANDLED")
                    self.parw.showTable()
                    def setFalse(): self.eventHandled=False
                    QTimer.singleShot(0, setFalse)
                return super().focusOutEvent(arg__1)
        ##CONTAINS PAIRS: COL_NAME(str) <=> FILTER_TEXT(LineEdit)
        self.filters = {str(header) : LineEdit("").init("Фильтр",self) for header in self.curHeaders}

        self.ui.addBtn.triggered.connect(self.addNewRow)
        self.ui.backBtn.triggered.connect(self.close)
        self.ui.saveBtn.setDisabled(True)
        self.ui.saveBtn_2.setDisabled(True)
        self.ui.reportBtn.triggered.connect(self.createReport)

        self.ui.tableWidget.customContextMenuRequested.connect(self.showContextMenu)

        self.showTable()
        self.updateStatus('OK')

    def showTable(self, sortingColName:str="", AZ:int=0):
        print([i.text() for i in self.filters.values()])
        finalRq = f"select distinct * from {self.tableName}"
        if not all([i.text() == "" or i.text().isspace() for i in self.filters.values()]): ##USING FILTERS
            print("UPDATING WITH FILTERS")
            filRq = []
            for col, lineedit in self.filters.items():
                if not (lineedit.text() == "" or lineedit.text().isspace()):
                    filRq += [f"{col}::text like \'%{lineedit.text()}%\'"]
            finalRq += ' where ' + ' and '.join(filRq)
        if AZ != 0:
            finalRq += f' order by {sortingColName}'
        if AZ == -1:
            finalRq += ' desc;'
        else:
            finalRq += ';'
        print(f'TRYING TO COMPLETE: {finalRq}')
        self.curTable = DBParser.execute(finalRq)
        if (len(self.curTable) == 0):
            print(f'NO ROWS HAVE BEEN FOUND')
            self.ACTIVE = False
            self.updateStatus("Подходящих строк не найдено! Пожалуйста очистите строку фильтров!")
        self.rowCount = len(self.curTable) + 1 #ADDING ONE ROW FOR FILTERS
        self.colCount = len(self.curTable[0]) if self.ACTIVE else self.colCount
        print(f'{self.rowCount - 1} ROWS HAVE BEEN FOUND')
        self.updateStatus(f"{self.rowCount - 1} записей найдено!")
        self.ui.tableWidget.setRowCount(self.rowCount)
        self.ui.tableWidget.setColumnCount(self.colCount)
        for col in range(self.colCount):
            self.ui.tableWidget.setCellWidget(0,col,self.filters[self.curHeaders[col]])
        for row in range(1,self.rowCount):
            for col in range(self.colCount):
                self.ui.tableWidget.setItem(row,col,QTableWidgetItem(str(self.curTable[row - 1][col])))
        self.ui.tableWidget.setHorizontalHeaderLabels([niceTableHeaders[self.tableName][i] for i in self.curHeaders])
    
    def addNewRow(self):
        if not self.ACTIVE: return
        def onFinish():
            wrap = lambda x: f"\'{x.text()}\'"
            print(DBParser.execute(f"insert into {self.tableName} values ({','.join(map(lambda x: wrap(x),self.dialog.cols.values()))}) returning *;"))
            self.showTable()
        self.dialog = AddRowDialog(self.curHeaders, self.tableName)
        self.dialog.accepted.connect(onFinish)
        self.dialog.show()

    def createReport(self):
        if not self.ACTIVE: return
        print("Starting report")
        self.dialog = SelectReportDialog(self)
        self.dialog.show()
        self.dialog.rejected.connect(lambda: self.show())
        self.hide()

    def showContextMenu(self, pos):
        if not self.ACTIVE: return
        rowInd = self.ui.tableWidget.indexAt(pos).row()
        colInd = self.ui.tableWidget.indexAt(pos).column()
        print(f"CONTEXT MANAGER REQUESTED ON CELL: {rowInd}, {colInd}")
        menu = QMenu()
        editAction = menu.addAction("Редактировать запись")
        deleteAction = menu.addAction("Удалить запись")
        sortAZAction = menu.addAction("Сортировать в алф порядке")
        sortZAAction = menu.addAction("Сортировать в обр алф порядке")
        action = menu.exec(self.ui.tableWidget.mapToGlobal(pos))
        if action == editAction:
            self.editRow(rowInd)
        elif action == deleteAction:
            self.deleteRow(rowInd)
        elif action == sortAZAction:
            self.sortCol(self.curHeaders[colInd], 1)
        elif action == sortZAAction:
            self.sortCol(self.curHeaders[colInd], -1)

    def editRow(self, rowInd:int):
        if not self.ACTIVE: return
        self.dialog = EditRowDialog({colName:curValue for colName,curValue in zip(self.curHeaders,[self.ui.tableWidget.item(rowInd,i).text() for i in range(self.colCount)])},self.tableName)
        id = self.ui.tableWidget.item(rowInd, 0).text()
        def onFinish():
            wrap = lambda x: f"{x[0]}=\'{x[1].text()}\'"
            print(DBParser.execute(f"update {self.tableName} set {','.join(map(lambda pair: wrap(pair),self.dialog.cols.items()))} where id=\'{id}\' returning *;"))
            self.showTable()
        self.dialog.accepted.connect(onFinish)
        self.dialog.show()

    def deleteRow(self, rowInd:int):
        if not self.ACTIVE: return
        id = self.ui.tableWidget.item(rowInd, 0).text()
        self.dialog = QMessageBox.warning(self, 'Удаление записи',"Вы уверены, что хотите удалить запись (ряд)?",QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if self.dialog == QMessageBox.StandardButton.Yes:
            print(DBParser.execute(f"delete from {self.tableName} where id = \'{id}\' returning *;"))
            self.showTable()
    
    def sortCol(self, sortingCol:str, AZ:int):
        if not self.ACTIVE: return
        self.showTable(sortingCol, AZ)

    def updateStatus(self, newStatus:str):
        print(newStatus)
        self.ui.statusbar.showMessage(newStatus)

    def closeEvent(self, event):
        self.parw.show()
        return super().closeEvent(event)

class SelectReportDialog(QDialog):
    def __init__(self, parw):
        super(SelectReportDialog, self).__init__()
        self.ui = Ui_selectReportView()
        self.ui.setupUi(self)
        self.parw = parw

        self.ui.listWidget.addItem("Список пациентов врача")
        self.ui.listWidget.addItem("Список свободных врачей на дату")
        self.ui.listWidget.addItem("Количество посещений пациентов")
        self.ui.listWidget.itemDoubleClicked.connect(self.clicked)

    def clicked(self, pos:QListWidgetItem):
        theme = pos.text()
        if theme == "Список пациентов врача":
            def show_choice_dialog(widget:QWidget):
                options = [i[0] for i in DBParser.execute("select full_name from doctors;")]
                print(options)
                choice, ok = QInputDialog.getItem(widget, "Выберите доктора:", "Выберите одного:", options, 0, False)
                if ok and choice:
                    self.reportWindow = reportViewWindow(theme,self.parw)
                    self.reportWindow.setData(choice)
                    self.reportWindow.show()
                    widget.close()

            self.widget = QWidget()
            self.button = QPushButton("Выберите доктора", self.widget)
            self.button.clicked.connect(lambda: show_choice_dialog(self.widget))
            self.button.resize(200, 40)
            self.button.move(50, 50)
            self.widget.setWindowModality(QtCore.Qt.WindowModality.ApplicationModal)
            self.widget.show()
        elif theme == "Список свободных врачей на дату":
            def onFinish(widget:QDialog):
                if self.datetime1.dateTime() >= self.datetime2.dateTime():
                    print("Wrong date")
                    selectDate()
                    return
                self.reportWindow = reportViewWindow(theme,self.parw)
                self.reportWindow.setData([self.datetime1.dateTime().toString("yyyy-MM-dd HH:mm:ss"),self.datetime2.dateTime().toString("yyyy-MM-dd HH:mm:ss")])
                self.reportWindow.show()
                widget.close()

            def selectDate():
                self.widget = QDialog()
                layout = QVBoxLayout()
                self.datetime1 = QDateTimeEdit(self.widget)
                self.datetime1.setDisplayFormat("yyyy-MM-dd HH:mm:ss")
                layout.addWidget(QLabel("Выберите начальный порог временного промежутка:"))
                layout.addWidget(self.datetime1)
                self.datetime2 = QDateTimeEdit(self.widget)
                self.datetime2.setDisplayFormat("yyyy-MM-dd HH:mm:ss")
                layout.addWidget(QLabel("Выберите конечный порог временного промежутка:"))
                layout.addWidget(self.datetime2)
                ok_button = QPushButton("OK")
                ok_button.clicked.connect(lambda: onFinish(self.widget))
                layout.addWidget(ok_button)
                self.widget.setLayout(layout)
                self.widget.setWindowModality(QtCore.Qt.WindowModality.ApplicationModal)
                self.widget.show()
            selectDate()
        elif theme == "Количество посещений пациентов":
            def selectData():
                self.widget = QDialog()
                layout = QVBoxLayout()
                res = [i[0] for i in DBParser.execute("select full_name from clients;")]
                self.options = {i:QCheckBox(i) for i in res}
                for option in self.options.values():
                    layout.addWidget(option)
                ok_button = QPushButton("OK")
                ok_button.clicked.connect(lambda: show_selected_options(self.widget))
                layout.addWidget(ok_button)
                self.widget.setLayout(layout)
                self.widget.setWindowModality(QtCore.Qt.WindowModality.ApplicationModal)
                self.widget.show()

            def show_selected_options(widget):
                selected_options = [name for name, checkbox in self.options.items() if checkbox.isChecked()]
                if selected_options:
                    self.reportWindow = reportViewWindow(theme, self.parw)
                    self.reportWindow.setData(selected_options)
                    self.reportWindow.show()
                else:
                    selectData()
                widget.close()
                
            selectData()

class reportViewWindow(QMainWindow):
    def __init__(self, theme:str, parw):
        super(reportViewWindow, self).__init__()
        self.ui = Ui_curTableView()
        self.ui.setupUi(self)
        self.baseRq = {
            "Список пациентов врача":"select doc.full_name as doctor, cl.full_name as client, vis.data_start as data_start, vis.data_end as data_end from doctors doc join (visits vis join clients cl on vis.client_id=cl.id) on doc.id = vis.doctor_id where doc.full_name like \'%{doctor}%\'",
            "Список свободных врачей на дату":"select doc.full_name as doctor,sp.name as spec from doctors doc full outer join visits vis on doc.id=vis.doctor_id left outer join specs sp on doc.spec_id=sp.id where (vis.data_start not between \'{data_start}\' and \'{data_end}\' and vis.data_end not between \'{data_start}\' and \'{data_end}\') or (vis.data_start is null and vis.data_end is null)",
            "Количество посещений пациентов":"select cl.full_name as client,count(vis.client_id) as visit_count from clients cl full outer join visits vis on cl.id=vis.client_id where {names} group by cl.full_name"
        }[theme]

        self.setWindowModality(QtCore.Qt.WindowModality.ApplicationModal)

        self.theme = theme
        self.ACTIVE = True

        self.tableName = theme
        # self.ui.tableNameLabel.setText(niceTableNames[theme])

        self.parw = parw

        self.curHeaders = {
            "Список пациентов врача":["doctor","client","data_start","data_end"],
            "Список свободных врачей на дату":["doctors", "spec"],
            "Количество посещений пациентов":["client","visit_count"]
        }[theme]

        print(self.curHeaders)
        self.ui.tableWidget.verticalHeader().setVisible(False)

        class LineEdit(QLineEdit):
            def init(self,arg_1,parw):
                self.setPlaceholderText(arg_1)
                self.parw = parw
                self.eventHandled = False
                return self
            def focusOutEvent(self, arg__1):
                print("FILTER CELL LOST FOCUS")
                if not self.eventHandled:
                    self.eventHandled = True
                    print("FILTER CELL DONT HAVE EVENT_HANDLED")
                    self.parw.showTable()
                    def setFalse(): self.eventHandled=False
                    QTimer.singleShot(0, setFalse)
                return super().focusOutEvent(arg__1)
        ##CONTAINS PAIRS: COL_NAME(str) <=> FILTER_TEXT(LineEdit)
        self.filters = {str(header) : LineEdit("").init("Фильтр",self) for header in self.curHeaders}

        self.ui.addBtn.setDisabled(True)
        self.ui.backBtn.setDisabled(True)
        self.ui.saveBtn.setDisabled(True)
        self.ui.saveBtn_2.setDisabled(True)
        self.ui.reportBtn.setDisabled(True)
        self.saveReportBtn = QAction(self)
        self.saveReportBtn.triggered.connect(self.saveReport)
        self.ui.menu.addAction(self.saveReportBtn)

        self.ui.tableWidget.customContextMenuRequested.connect(self.showContextMenu)
    
    def saveReport(self):
        data = []
        for row in range(1,self.ui.tableWidget.rowCount()):
            row_data = []
            for column in range(self.ui.tableWidget.columnCount()):
                item = self.ui.tableWidget.item(row, column)
                row_data.append(item.text() if item else None)
            data.append(row_data)
        df = pd.DataFrame(data)
        print(df)

        doc = Document()
        table = doc.add_table(rows = df.shape[0] + 1, cols=df.shape[1])
        for j in range(df.shape[1]):
            table.cell(0,j).text = self.curHeaders[j]
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                table.cell(i + 1, j).text = str(df.iat[i, j])
        doc.save('lastreport.docx')

    def showTable(self, sortingColName:str="", AZ:int=0):
        print([i.text() for i in self.filters.values()])
        finalRq = f"select * from ({self.baseRq}) as temptable"
        if not all([i.text() == "" or i.text().isspace() for i in self.filters.values()]): ##USING FILTERS
            print("UPDATING WITH FILTERS")
            filRq = []
            for col, lineedit in self.filters.items():
                if not (lineedit.text() == "" or lineedit.text().isspace()):
                    filRq += [f"temptable.{col}::text like \'%{lineedit.text()}%\'"]
            finalRq += ' where ' + ' and '.join(filRq)
        if AZ != 0:
            finalRq += f' order by temptable.{sortingColName}'
        if AZ == -1:
            finalRq += ' desc;'
        else:
            finalRq += ';'
        print(f'TRYING TO COMPLETE: {finalRq}')
        self.curTable = DBParser.execute(finalRq)
        if (len(self.curTable) == 0):
            print(f'NO ROWS HAVE BEEN FOUND')
            self.ACTIVE = False
        self.rowCount = len(self.curTable) + 1 #ADDING ONE ROW FOR FILTERS
        self.colCount = len(self.curTable[0]) if self.ACTIVE else len(self.curHeaders)
        print(f'{self.rowCount - 1} ROWS HAVE BEEN FOUND')
        self.ui.tableWidget.setRowCount(self.rowCount)
        self.ui.tableWidget.setColumnCount(self.colCount)
        for col in range(self.colCount):
            self.ui.tableWidget.setCellWidget(0,col,self.filters[self.curHeaders[col]])
        for row in range(1,self.rowCount):
            for col in range(self.colCount):
                self.ui.tableWidget.setItem(row,col,QTableWidgetItem(str(self.curTable[row - 1][col])))
        self.ui.tableWidget.setHorizontalHeaderLabels([i for i in self.curHeaders])
    
    def showContextMenu(self, pos):
        if not self.ACTIVE: return
        rowInd = self.ui.tableWidget.indexAt(pos).row()
        colInd = self.ui.tableWidget.indexAt(pos).column()
        print(f"CONTEXT MANAGER REQUESTED ON CELL: {rowInd}, {colInd}")
        menu = QMenu()
        saveAction = menu.addAction("Сохранить отчет")
        sortAZAction = menu.addAction("Сортировать в алф порядке")
        sortZAAction = menu.addAction("Сортировать в обр алф порядке")
        action = menu.exec(self.ui.tableWidget.mapToGlobal(pos))
        if action == saveAction:
            self.saveReport()
        elif action == sortAZAction:
            self.sortCol(self.curHeaders[colInd], 1)
        elif action == sortZAAction:
            self.sortCol(self.curHeaders[colInd], -1)

    def sortCol(self, sortingCol:str, AZ:int):
        if not self.ACTIVE: return
        self.showTable(sortingCol, AZ)

    def setData(self, data):
        self.data = data
        if self.theme == "Список пациентов врача":
            self.baseRq = self.baseRq.format(doctor=data)
        elif self.theme == "Список свободных врачей на дату":
            self.baseRq = self.baseRq.format(data_start=data[0], data_end=data[1])
        else:
            s = ' or '.join([f"cl.full_name=\'{i}\'" for i in self.data])
            self.baseRq = self.baseRq.format(names=s)
        self.showTable()